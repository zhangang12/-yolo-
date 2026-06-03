#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
审查报告导出（Word）  report_docx.py
=================================================
把 e2e_demo 的产出（e2e_structured.json + e2e_annotated.png）汇编成一份
可交付的 Word 审查报告：封面信息 → 结论汇总 → 逐条检查表 → 标注图 → 免责说明。

用法：
  python report_docx.py  e2e输出目录  [--out 报告.docx]
  # 目录里需有 e2e_structured.json（必需）与 e2e_annotated.png（可选，有则插图）

依赖：python-docx
"""
import sys, os, json, argparse

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass


def _set_cell(cell, text, bold=False, color=None, align=None):
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def build(out_dir, out_path=None):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    js = os.path.join(out_dir, "e2e_structured.json")
    if not os.path.exists(js):
        print(f"❌ 找不到 {js}（请先在端到端预审里跑一遍）"); sys.exit(1)
    with open(js, encoding="utf-8") as f:
        data = json.load(f)
    findings = data.get("findings", [])
    img = os.path.join(out_dir, "e2e_annotated.png")
    out_path = out_path or os.path.join(out_dir, "审查报告.docx")

    doc = Document()
    # 默认中文字体
    doc.styles["Normal"].font.name = "微软雅黑"
    doc.styles["Normal"].font.size = Pt(10.5)

    # ---- 标题 ----
    t = doc.add_heading("消防设计图纸 AI 预审报告", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- 基本信息 ----
    doc.add_heading("一、基本信息", level=1)
    info = doc.add_table(rows=0, cols=2); info.style = "Light Grid Accent 1"
    for k, v in [("图纸", data.get("drawing", "—")),
                 ("数据来源", "图纸文字层直读（设计院声称值）" if data.get("source") == "vector_text"
                  else "矢量提取 + OCR 识别（含噪声）"),
                 ("检查项总数", str(len(findings)))]:
        r = info.add_row().cells
        _set_cell(r[0], k, bold=True); _set_cell(r[1], v)

    # ---- 结论汇总 ----
    n_fail = sum(1 for f in findings if f.get("passed") is False)
    n_rev = sum(1 for f in findings if f.get("review_required"))
    n_pass = sum(1 for f in findings if f.get("passed") is True)
    doc.add_heading("二、结论汇总", level=1)
    verdict = "不通过（存在不合规项）" if n_fail else ("基本通过（含待人工复核项）" if n_rev else "通过")
    p = doc.add_paragraph()
    p.add_run("审查结论：").bold = True
    rr = p.add_run(verdict)
    rr.bold = True
    rr.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B) if n_fail else (
        RGBColor(0xE8, 0x59, 0x0C) if n_rev else RGBColor(0x2B, 0x8A, 0x3E))
    doc.add_paragraph(f"通过 {n_pass} 项 ／ 不合规 {n_fail} 项 ／ 待人工复核 {n_rev} 项。")

    # ---- 逐条检查 ----
    doc.add_heading("三、逐条检查明细", level=1)
    tbl = doc.add_table(rows=1, cols=5); tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["对象", "实测/声称值", "规范要求", "结论", "规范出处"]):
        _set_cell(hdr[i], h, bold=True, align="center")
    for f in findings:
        passed = f.get("passed")
        if passed is False:
            concl, col = "不通过", (0xC0, 0x39, 0x2B)
        elif f.get("review_required"):
            concl, col = "待人工复核", (0xE8, 0x59, 0x0C)
        else:
            concl, col = "通过", (0x2B, 0x8A, 0x3E)
        c = tbl.add_row().cells
        _set_cell(c[0], f.get("target_id") or f.get("category", "—"))
        _set_cell(c[1], f.get("value", "—"))
        thr = f.get("threshold")
        _set_cell(c[2], f"≤{thr}" if thr is not None else "—")
        _set_cell(c[3], concl, bold=True, color=col)
        _set_cell(c[4], f.get("source", "—"))

    # ---- 标注图 ----
    if os.path.exists(img):
        doc.add_heading("四、原图标注", level=1)
        doc.add_paragraph("绿框=合规，红框=超限，黄框=待人工复核。")
        try:
            doc.add_picture(img, width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"（标注图插入失败：{e}）")

    # ---- 自洽检查 / 免责 ----
    doc.add_heading("五、说明与免责", level=1)
    for c in data.get("consistency", []):
        doc.add_paragraph(c, style="List Bullet")
    doc.add_paragraph(
        "本报告由 AI 预审系统自动生成，规则判定基于确定性规则引擎并附规范出处。"
        "其中面积等数值若来自图纸文字层，为设计院标注的声称值，未经几何复核；"
        "涉及空间关系（出口个数/间距、净宽、门开启方向等）的指标需结合人工审查。"
        "本报告供审查参考，不替代法定审查结论。", style="Intense Quote")

    doc.save(out_path)
    print(f"[报告] 已生成 Word 报告：{out_path}")
    print(f"[报告] 结论：{verdict}（通过{n_pass}/不合规{n_fail}/待复核{n_rev}）")
    return out_path


def main():
    ap = argparse.ArgumentParser(description="e2e 产出 → Word 审查报告")
    ap.add_argument("out_dir", help="e2e 输出目录（含 e2e_structured.json）")
    ap.add_argument("--out", help="报告输出路径(.docx)")
    a = ap.parse_args()
    build(a.out_dir, a.out)


if __name__ == "__main__":
    main()
