#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MVP Word 审查报告  mvp_report_docx.py
=================================================
仿照真实"消防设计文件审查意见表"格式(参考 J2025-004 三份样本),
把 mvp_e2e 的 findings.json 整理成可交付的 .docx,包含:

  - 头部基本信息表(工程名称 / 建设单位 / 设计单位 / 审查阶段 等可填)
  - 结论汇总(总数 / 通过 / 不通过 / 待复核)
  - 主表:序号 | 图号 | 审查意见(中文,带规范出处) | AI 判定值 | 设计回复 | 复核情况
  - 嵌入标注图
  - 免责说明

用法:
  python mvp_report_docx.py  <findings.json>  <annotated.jpg>  <out.docx>
        [--station-name 嘉宾站]
        [--design-stage 初审]
        [--reviewer 'AI 预审系统']
"""
import os, sys, json, argparse
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- 中文规则名映射(rule_id → 人话审查意见模板) ----------
# 与 rules/rules.json 的 name 字段一致,这里再加"中文化模板"用于审查报告
RULE_HUMAN = {
    "WIDTH-DOOR-EVAC-001":
        "{tid} 疏散门净宽 {value}m,不满足规范要求 ≥{threshold}m,应加大门洞宽度。",
    "DOOR-SWING-EVAC-001":
        "{tid} 疏散门未向疏散方向开启,不满足规范要求(向疏散方向开启),应调整门扇方向。",
    "DOOR-CLASS-FIREWALL-001":
        "{tid} 防火墙上防火门等级不是甲级,应改为甲级防火门。",
    "DOOR-CLASS-PROJECT-ALL-A-001":
        "{tid} 本工程承诺所有防火门均为甲级,此门等级不符,应修改。",
    "AREA-PUB-001":
        "{tid} 站厅公共区防火分区面积 {value}㎡,超规范上限 {threshold}㎡,应增设防火分隔。",
    "AREA-EQUIP-UG-001":
        "{tid} 地下设备区每防火分区面积 {value}㎡,超规范上限 {threshold}㎡。",
    "AREA-EQUIP-AG-001":
        "{tid} 地上设备区每防火分区面积 {value}㎡,超规范上限 {threshold}㎡。",
    "EXIT-NUM-PUB-001":
        "{tid} 公共区安全出口仅 {value} 个,不满足规范要求 ≥{threshold} 个,应增设。",
    "EXIT-NUM-EQUIP-001":
        "{tid} 设备区防火分区安全出口 {value} 个,不满足规范要求 ≥{threshold} 个(含相邻分区共用防火墙上的常开甲级防火门)。",
    "EXIT-NUM-EQUIP-STAFFED-001":
        "{tid} 有人值守的设备区防火分区,直通地面的安全出口 {value} 个,不满足规范要求 ≥{threshold} 个。",
    "EXIT-DIST-SAME-DIR-001":
        "{tid} 同方向两安全出口距离 {value}m,不满足规范要求 ≥{threshold}m。",
    "EXIT-DIST-ADJ-001":
        "{tid} 相邻两安全出口距离 {value}m,不满足规范要求 ≥{threshold}m。",
    "EVAC-DIST-ANY-001":
        "{tid} 任一点至安全出口疏散距离 {value}m,超规范上限 {threshold}m,应增设安全出口或调整布局。",
    "EVAC-DIST-EQUIP-BETWEEN-001":
        "{tid} 设备区房间门至安全出口距离 {value}m(位于两个出口之间),超规范上限 {threshold}m。",
    "EVAC-DIST-EQUIP-DEADEND-001":
        "{tid} 设备区房间门至安全出口距离 {value}m(袋形走道尽端),超规范上限 {threshold}m。",
    "WIDTH-CORR-EVAC-001":
        "{tid} 疏散走道净宽 {value}m,不满足规范要求 ≥{threshold}m。",
    "WIDTH-STAIR-EVAC-001":
        "{tid} 疏散楼梯净宽 {value}m,不满足规范要求 ≥{threshold}m。",
    "WIDTH-EQUIP-CORR-SINGLE-001":
        "{tid} 设备区单面布置走道净宽 {value}m,不满足规范要求 ≥{threshold}m。",
    "WIDTH-EQUIP-CORR-DOUBLE-001":
        "{tid} 设备区双面布置走道净宽 {value}m,不满足规范要求 ≥{threshold}m。",
    "LANDING-WIDTH-001":
        "{tid} 出地面楼梯平台净宽 {value}m,不满足规范要求 ≥{threshold}m。",
    "SHOP-COUNT-PUB-001":
        "公共区商铺数量 {value} 个,超规范上限 {threshold} 个。",
    "SHOP-AREA-EACH-001":
        "{tid} 单个商铺面积 {value}㎡,超规范上限 {threshold}㎡。",
    "SHOP-AREA-TOTAL-PUB-001":
        "公共区商铺总面积 {value}㎡,超规范上限 {threshold}㎡。",
    "SHOP-OPENING-DIST-001":
        "{tid} 商铺间洞口距离 {value}m,不满足规范要求 ≥{threshold}m。",
    "SHUTTER-WIDTH-LE30-001":
        "{tid} 防火卷帘宽度 {value}m,超规范上限(洞口宽/3 且 ≤20m)。",
    "SHUTTER-WIDTH-GT30-001":
        "{tid} 洞口宽 >30m 的防火卷帘宽度 {value}m,不满足规范要求。",
    "FIRE-CLEARANCE-MULTI-001":
        "{tid}({target.building_name}) 出入口/风亭与周边多层民用建筑(一二级)间距 {value}m,不满足规范 ≥{threshold}m。",
    "FIRE-CLEARANCE-HIGHRISE-001":
        "{tid}({target.building_name}) 出入口/风亭与周边高层民用建筑间距 {value}m,不满足规范 ≥{threshold}m。",
    "FIRE-CLEARANCE-GASSTATION-001":
        "{tid}({target.building_name}) 出入口与加油加气加氢站安全间距 {value}m,不满足规范 ≥{threshold}m。",
    "VENT-FRESH-DIST-001":
        "{tid} 新风口与排风/活塞风口距离 {value}m,不满足规范 ≥{threshold}m。",
    "VENT-HEIGHT-001":
        "{tid} 排风/活塞口未高于新风口,不满足规范要求。",
    "CORR-UG-ENTRANCE-LEN-001":
        "{tid} 地下出入口通道长度 {value}m,超规范上限 {threshold}m。",
}


def _human_message(f):
    """根据 finding 产出中文审查意见。"""
    tpl = RULE_HUMAN.get(f["rule_id"])
    if not tpl:
        # fallback: 原 message + 中文修饰
        return f.get("message", "")
    try:
        return tpl.format(tid=f.get("target_id", ""),
                          value=f.get("value", ""),
                          threshold=f.get("threshold", ""))
    except Exception:
        return f.get("message", "")


def _set_cell(cell, text, bold=False, align=None, font_size=None, color=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.text = ""
    run = p.add_run(text)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_col_widths(table, widths_cm):
    """精准设置每列宽度(cm)。"""
    for i, w in enumerate(widths_cm):
        for row in table.rows:
            row.cells[i].width = Cm(w)


def _shade_cell(cell, hex_color):
    """给单元格涂色(底纹)。hex_color 如 'FFE5E5'。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def build(findings_path, annotated_img, out_docx,
          station_name="未指定", design_stage="初审",
          reviewer="AI 预审系统", source_pdf_name=None):
    findings = json.load(open(findings_path, encoding="utf-8"))

    fails = [f for f in findings if f.get("passed") is False]
    reviews = [f for f in findings if f.get("review_required")]
    passes = [f for f in findings if f.get("passed") is True]

    doc = Document()
    # 全文默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(10.5)

    # ---------- 标题 ----------
    title = doc.add_heading("消防设计图纸 AI 预审意见表", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("(预审系统自动生成,供审查参考,不替代法定审查结论)")
    sr.font.size = Pt(9); sr.italic = True

    # ---------- 基本信息表(模仿 J2025-004 头部) ----------
    info = doc.add_table(rows=4, cols=4)
    info.style = "Light Grid Accent 1"
    _set_col_widths(info, [3, 6, 3, 6])

    _set_cell(info.cell(0, 0), "工程名称", bold=True)
    _set_cell(info.cell(0, 1), station_name)
    _set_cell(info.cell(0, 2), "审查阶段", bold=True)
    _set_cell(info.cell(0, 3), design_stage)

    _set_cell(info.cell(1, 0), "图纸文件", bold=True)
    _set_cell(info.cell(1, 1), source_pdf_name or "—")
    _set_cell(info.cell(1, 2), "审查日期", bold=True)
    _set_cell(info.cell(1, 3), datetime.now().strftime("%Y.%m.%d"))

    _set_cell(info.cell(2, 0), "审查人(系统)", bold=True)
    _set_cell(info.cell(2, 1), reviewer)
    _set_cell(info.cell(2, 2), "规则引擎版本", bold=True)
    _set_cell(info.cell(2, 3), "rules.json (37 条 / 8 大类)")

    _set_cell(info.cell(3, 0), "数据来源", bold=True)
    _set_cell(info.cell(3, 1), "CVAT 标注真值 + 矢量 PDF 几何量测")
    _set_cell(info.cell(3, 2), "审查范围", bold=True)
    _set_cell(info.cell(3, 3), "防火分区/疏散/安全出口/防火门/防火间距等")

    doc.add_paragraph()

    # ---------- 结论汇总 ----------
    doc.add_heading("一、审查结论", level=1)
    p = doc.add_paragraph()
    p.add_run("总检查项:").bold = True
    p.add_run(f"{len(findings)} 项    ")
    p.add_run("通过:").bold = True
    rp = p.add_run(f"{len(passes)} 项    ")
    rp.font.color.rgb = RGBColor(0x2B, 0x8A, 0x3E)
    p.add_run("不合规:").bold = True
    rf = p.add_run(f"{len(fails)} 项    ")
    rf.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p.add_run("待人工复核:").bold = True
    rr = p.add_run(f"{len(reviews)} 项")
    rr.font.color.rgb = RGBColor(0xC0, 0x83, 0x2B)

    conclude = doc.add_paragraph()
    if len(fails) == 0 and len(reviews) == 0:
        cr = conclude.add_run("初步结论:全部通过(无强条不合规)")
        cr.font.color.rgb = RGBColor(0x2B, 0x8A, 0x3E); cr.bold = True
    elif len(fails) == 0:
        cr = conclude.add_run(f"初步结论:无强条不合规,但有 {len(reviews)} 项需人工复核")
        cr.font.color.rgb = RGBColor(0xC0, 0x83, 0x2B); cr.bold = True
    else:
        cr = conclude.add_run(f"初步结论:不通过(存在 {len(fails)} 项强条不合规,需修改)")
        cr.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B); cr.bold = True

    # ---------- 主表(按规则分类聚合) ----------
    doc.add_heading("二、不合规明细", level=1)
    if not fails:
        doc.add_paragraph("(无)")
    else:
        from collections import defaultdict
        by_rule = defaultdict(list)
        for f in fails:
            by_rule[f["rule_id"]].append(f)
        # 主表表头与真实审查报告一致
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["序号", "图上编号", "审查意见", "规范出处", "设计回复", "复核情况"]):
            _set_cell(hdr[i], h, bold=True, align="center")
            _shade_cell(hdr[i], "D4E6F1")
        _set_col_widths(tbl, [1.2, 1.8, 7.5, 4, 2.5, 2])
        seq = 0
        for rid, items in sorted(by_rule.items(), key=lambda x: -len(x[1])):
            for f in items:
                seq += 1
                row = tbl.add_row().cells
                _set_cell(row[0], str(seq), align="center")
                _set_cell(row[1], f.get("target_id", "—"), align="center")
                _set_cell(row[2], _human_message(f))
                _set_cell(row[3], f.get("source", "—"))
                _set_cell(row[4], "")  # 设计回复(空,留给人工)
                _set_cell(row[5], "□通过  □不通过", align="center")

    # ---------- 待复核 ----------
    if reviews:
        doc.add_heading(f"三、待人工复核({len(reviews)} 项)", level=1)
        doc.add_paragraph(
            "以下检查项因数据不全(如分区面积未标、几何关系不明等)无法自动判定,需人工复核:"
        )
        tbl2 = doc.add_table(rows=1, cols=4)
        tbl2.style = "Light Grid Accent 1"
        hdr2 = tbl2.rows[0].cells
        for i, h in enumerate(["序号", "图上编号", "复核内容", "规范出处"]):
            _set_cell(hdr2[i], h, bold=True, align="center")
            _shade_cell(hdr2[i], "FCF3CF")
        _set_col_widths(tbl2, [1.2, 2, 10, 5])
        for i, f in enumerate(reviews[:30], 1):    # 最多列 30 条
            row = tbl2.add_row().cells
            _set_cell(row[0], str(i), align="center")
            _set_cell(row[1], f.get("target_id", "—"), align="center")
            _set_cell(row[2], f.get("message", ""))
            _set_cell(row[3], f.get("source", "—"))

    # ---------- 标注图 ----------
    doc.add_heading("四、违规位置标注图", level=1)
    if annotated_img and os.path.exists(annotated_img):
        doc.add_paragraph("红框 = 不合规(对应「不合规明细」表的图上编号),旁边附 ID 与规则编号:")
        doc.add_picture(annotated_img, width=Cm(16))
    else:
        doc.add_paragraph("(标注图未生成或不可读)")

    # ---------- 免责 ----------
    doc.add_heading("五、说明与免责", level=1)
    p = doc.add_paragraph()
    p.add_run("本报告由 AI 预审系统自动生成。").bold = True
    doc.add_paragraph(
        "面积等数值若来自图纸文字层,为设计院的声称值,未经几何复核;"
        "涉及空间关系(出口数 / 间距 / 净宽 / 门开启方向等)的判定依据 CVAT 真值标注 + 矢量 PDF 几何量测,"
        "结论受标注完整度影响。本报告供审查参考,不替代法定审查结论。"
    )
    doc.add_paragraph(
        f"规则引擎采用确定性 if-else 判定,每条违规均附规范出处(GB 50016 / GB 50157 / GB 51298 等)。"
    )

    doc.save(out_docx)
    print(f"[完成] 报告: {out_docx}")
    print(f"  不合规 {len(fails)} 项 / 待复核 {len(reviews)} 项 / 通过 {len(passes)} 项")
    return out_docx


def main():
    ap = argparse.ArgumentParser(description="MVP 端到端 → 仿真实审查意见表 Word 报告")
    ap.add_argument("findings", help="mvp_e2e 输出的 *_findings.json")
    ap.add_argument("annotated", help="mvp_e2e 输出的 *_e2e_fail.jpg")
    ap.add_argument("out", help="输出 docx 路径")
    ap.add_argument("--station-name", default="未指定")
    ap.add_argument("--design-stage", default="初审")
    ap.add_argument("--reviewer", default="AI 预审系统 (本项目 MVP)")
    ap.add_argument("--source-pdf-name", default=None)
    a = ap.parse_args()
    build(a.findings, a.annotated, a.out, a.station_name, a.design_stage,
          a.reviewer, a.source_pdf_name)


if __name__ == "__main__":
    main()
